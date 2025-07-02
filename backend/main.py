from fastapi import FastAPI, HTTPException, UploadFile, File, Depends
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional
from openai import OpenAI
import os
from dotenv import load_dotenv
from sqlalchemy.orm import Session
from sqlalchemy import create_engine
import uuid
import json

# Загружаем переменные окружения с автоматическим определением среды
from load_env import load_environment
load_environment()

# Импорты для работы с базой данных
from database import get_db, engine, Base
from models import User, Session as DBSession, Resume, Analysis
from encryption import anonymize_for_ai, encrypt_resume_for_storage, decrypt_resume_from_storage, data_protector
from datetime import datetime
import PyPDF2
from docx import Document
import io
import pdfplumber
import fitz  # PyMuPDF

app = FastAPI(title="Career Mini App API", version="1.0.0")

# Создаем таблицы в базе данных
Base.metadata.create_all(bind=engine)  # Включено после настройки PostgreSQL на Beget

# CORS middleware для фронтенда
allowed_origins = os.getenv("ALLOWED_ORIGINS", "http://localhost:3000,https://expaai.github.io").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Health check endpoint для мониторинга
@app.get("/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

# Модели данных
class CareerAdviceRequest(BaseModel):
    user_goal: str
    experience_level: str
    current_role: Optional[str] = None
    interests: Optional[List[str]] = None

class CareerAdviceResponse(BaseModel):
    advice: str
    recommendations: List[str]
    next_steps: List[str]
    timestamp: datetime

class ResumeAnalysisRequest(BaseModel):
    resume_text: str
    target_position: Optional[str] = None

class ResumeAnalysisResponse(BaseModel):
    score: int  # 1-100
    strengths: List[str]
    improvements: List[str]
    missing_skills: List[str]
    recommendations: List[str]

class CareerTestRequest(BaseModel):
    answers: List[int]  # Ответы на вопросы теста
    test_type: str  # "personality", "skills", "interests"

class CareerTestResponse(BaseModel):
    result_type: str
    description: str
    recommended_careers: List[str]
    development_areas: List[str]

# Новые модели для логирования
class ProfessionSelectionRequest(BaseModel):
    user_id: Optional[str] = None
    selected_profession: str
    user_role: Optional[str] = None  # 'student' | 'professional'
    user_goal: Optional[str] = None
    timestamp: Optional[datetime] = None

class ProfessionSelectionResponse(BaseModel):
    success: bool
    message: str
    session_id: str

class JobMatchingRequest(BaseModel):
    session_id: str
    profession: str
    user_role: Optional[str] = None
    user_goal: Optional[str] = None
    job_url: Optional[str] = None

class JobMatchingResponse(BaseModel):
    jobs: List[dict]
    total_count: int
    session_id: str

class ResumeUploadResponse(BaseModel):
    success: bool
    extracted_text: str
    file_type: str
    message: str

class ResumeAnalysisAIRequest(BaseModel):
    resume_text: str
    profession: str
    job_url: Optional[str] = None

class ResumeAnalysisAIResponse(BaseModel):
    analysis: str
    success: bool
    message: str

# Новые модели для управления ролями и состоянием
class UserRoleRequest(BaseModel):
    user_id: str
    user_role: str  # 'student' | 'professional'
    timestamp: Optional[datetime] = None

class UserRoleResponse(BaseModel):
    success: bool
    message: str
    session_id: str

class SessionStateRequest(BaseModel):
    user_id: str

class SessionStateResponse(BaseModel):
    success: bool
    has_role: bool
    user_role: Optional[str] = None
    selected_profession: Optional[str] = None
    career_goal: Optional[str] = None
    session_id: Optional[str] = None
    current_screen: str  # 'role', 'goals', 'profession', 'advice', 'resume', 'main'

# Функции для обработки файлов
def extract_text_from_pdf(file_content: bytes) -> str:
    """Извлечение текста из PDF файла с несколькими методами"""
    
    # Метод 1: PyMuPDF (fitz) - самый надежный
    try:
        pdf_document = fitz.open(stream=file_content, filetype="pdf")
        text = ""
        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            text += page.get_text() + "\n"
        pdf_document.close()
        
        if text.strip():
            print(f"✅ PDF обработан через PyMuPDF: {len(text)} символов")
            return text.strip()
    except Exception as e:
        print(f"⚠️ PyMuPDF не смог обработать PDF: {str(e)}")
    
    # Метод 2: pdfplumber - хорош для таблиц и сложной разметки
    try:
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if text.strip():
            print(f"✅ PDF обработан через pdfplumber: {len(text)} символов")
            return text.strip()
    except Exception as e:
        print(f"⚠️ pdfplumber не смог обработать PDF: {str(e)}")
    
    # Метод 3: PyPDF2 - резервный вариант
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        
        if text.strip():
            print(f"✅ PDF обработан через PyPDF2: {len(text)} символов")
            return text.strip()
    except Exception as e:
        print(f"⚠️ PyPDF2 не смог обработать PDF: {str(e)}")
    
    # Если все методы не сработали
    raise ValueError("PDF файл не содержит извлекаемого текста или поврежден. Попробуйте сохранить файл в другом формате (DOCX или TXT) или используйте другой PDF-файл.")

def extract_text_from_docx(file_content: bytes) -> str:
    """Извлечение текста из DOCX файла"""
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        text = ""
        # Извлекаем текст из параграфов
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        if not text.strip():
            raise ValueError("DOCX файл не содержит текста")
        
        print(f"✅ DOCX обработан: {len(text)} символов")
        return text.strip()
    except Exception as e:
        print(f"⚠️ Ошибка при чтении DOCX: {str(e)}")
        raise ValueError(f"Не удалось обработать DOCX файл: {str(e)}. Убедитесь, что файл не поврежден.")

def extract_text_from_txt(file_content: bytes) -> str:
    """Извлечение текста из TXT файла с поддержкой различных кодировок"""
    encodings = ['utf-8', 'cp1251', 'latin-1', 'utf-16', 'ascii']
    
    for encoding in encodings:
        try:
            text = file_content.decode(encoding)
            if text.strip():
                print(f"✅ TXT обработан с кодировкой {encoding}: {len(text)} символов")
                return text.strip()
        except UnicodeDecodeError:
            continue
        except Exception as e:
            print(f"⚠️ Ошибка с кодировкой {encoding}: {str(e)}")
            continue
    
    # Если ни одна кодировка не подошла
    raise ValueError("Не удалось определить кодировку TXT файла. Попробуйте сохранить файл в UTF-8.")

# Эндпоинты
@app.get("/")
async def root():
    return {"message": "Career Mini App API", "version": "1.0.0"}

@app.post("/career-advice", response_model=CareerAdviceResponse)
async def get_career_advice(request: CareerAdviceRequest):
    """
    Генерация персонализированных карьерных советов на основе целей и опыта пользователя
    """
    try:
        # Формируем промпт для AI
        prompt = f"""
        Ты карьерный консультант. Дай персонализированный совет пользователю:
        
        Цель: {request.user_goal}
        Уровень опыта: {request.experience_level}
        Текущая роль: {request.current_role or 'Не указана'}
        Интересы: {', '.join(request.interests) if request.interests else 'Не указаны'}
        
        Предоставь:
        1. Общий совет (2-3 предложения)
        2. 3-5 конкретных рекомендаций
        3. 3-4 следующих шага
        
        Отвечай на русском языке, будь конкретным и практичным.
        """
        
        # Пока что возвращаем мок-данные (позже интегрируем с OpenAI)
        advice = f"Для достижения цели '{request.user_goal}' с уровнем опыта '{request.experience_level}' рекомендую сосредоточиться на развитии ключевых навыков и построении профессиональной сети."
        
        recommendations = [
            "Изучите актуальные технологии в вашей области",
            "Найдите ментора или карьерного консультанта",
            "Развивайте soft skills: коммуникация и лидерство",
            "Создайте план профессионального развития на год",
            "Участвуйте в профессиональных мероприятиях и конференциях"
        ]
        
        next_steps = [
            "Обновите резюме и LinkedIn профиль",
            "Определите 3 ключевых навыка для развития",
            "Найдите 5 компаний мечты и изучите их требования",
            "Запланируйте еженедельное время для обучения"
        ]
        
        return CareerAdviceResponse(
            advice=advice,
            recommendations=recommendations,
            next_steps=next_steps,
            timestamp=datetime.now()
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка генерации совета: {str(e)}")

@app.post("/resume-analysis", response_model=ResumeAnalysisResponse)
async def analyze_resume(request: ResumeAnalysisRequest):
    """
    Анализ резюме и предоставление рекомендаций по улучшению
    """
    try:
        # Мок-анализ резюме (позже интегрируем с AI)
        score = 75  # Базовая оценка
        
        strengths = [
            "Четкая структура резюме",
            "Релевантный опыт работы",
            "Указаны конкретные достижения"
        ]
        
        improvements = [
            "Добавить больше количественных метрик",
            "Улучшить описание ключевых навыков",
            "Добавить раздел с проектами"
        ]
        
        missing_skills = [
            "Современные технологии в области",
            "Навыки управления проектами",
            "Знание английского языка"
        ]
        
        recommendations = [
            "Переформулируйте достижения в формате STAR",
            "Добавьте ключевые слова из описания вакансии",
            "Сократите резюме до 1-2 страниц",
            "Добавьте ссылки на портфолио или проекты"
        ]
        
        return ResumeAnalysisResponse(
            score=score,
            strengths=strengths,
            improvements=improvements,
            missing_skills=missing_skills,
            recommendations=recommendations
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка анализа резюме: {str(e)}")

@app.post("/career-test", response_model=CareerTestResponse)
async def process_career_test(request: CareerTestRequest):
    """
    Обработка результатов карьерного теста
    """
    try:
        # Простая логика обработки теста (позже можно усложнить)
        avg_score = sum(request.answers) / len(request.answers)
        
        if request.test_type == "personality":
            if avg_score >= 4:
                result_type = "Лидер и новатор"
                description = "Вы прирожденный лидер, который любит инновации и изменения"
                recommended_careers = ["Менеджер проекта", "Предприниматель", "Консультант"]
            elif avg_score >= 3:
                result_type = "Командный игрок"
                description = "Вы отлично работаете в команде и цените стабильность"
                recommended_careers = ["Аналитик", "Специалист по продукту", "HR-менеджер"]
            else:
                result_type = "Независимый специалист"
                description = "Вы предпочитаете работать самостоятельно и глубоко погружаться в задачи"
                recommended_careers = ["Разработчик", "Дизайнер", "Исследователь"]
        else:
            result_type = "Универсальный профиль"
            description = "У вас разносторонние интересы и способности"
            recommended_careers = ["Продакт-менеджер", "Бизнес-аналитик", "Маркетолог"]
        
        development_areas = [
            "Развитие коммуникативных навыков",
            "Изучение новых технологий",
            "Развитие лидерских качеств",
            "Улучшение аналитических способностей"
        ]
        
        return CareerTestResponse(
            result_type=result_type,
            description=description,
            recommended_careers=recommended_careers,
            development_areas=development_areas
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка обработки теста: {str(e)}")

@app.post("/log-profession-selection", response_model=ProfessionSelectionResponse)
async def log_profession_selection(request: ProfessionSelectionRequest, db: Session = Depends(get_db)):
    """
    Логирование выбора профессии пользователем
    """
    try:
        # Создаем или получаем пользователя
        user = db.query(User).filter(User.telegram_id == request.user_id).first()
        if not user:
            user = User(
                telegram_id=request.user_id,
                username=f"user_{request.user_id}",
                data_processing_consent=True
            )
            db.add(user)
            db.commit()
            db.refresh(user)
        
        # Находим последнюю сессию пользователя или создаем новую
        db_session = db.query(DBSession).filter(
            DBSession.user_id == user.id
        ).order_by(DBSession.created_at.desc()).first()
        
        if db_session:
            # Обновляем существующую сессию
            db_session.selected_profession = request.selected_profession
            if request.user_role:
                db_session.career_stage = request.user_role
            if request.user_goal:
                db_session.career_goal = request.user_goal
            session_id = db_session.session_id
        else:
            # Создаем новую сессию
            session_id = str(uuid.uuid4())
            db_session = DBSession(
                session_id=session_id,
                user_id=user.id,
                selected_profession=request.selected_profession,
                career_stage=request.user_role,
                career_goal=request.user_goal
            )
            db.add(db_session)
        
        db.commit()
        
        # Логируем данные для отладки
        log_data = {
            "session_id": session_id,
            "user_id": request.user_id,
            "selected_profession": request.selected_profession,
            "user_role": request.user_role,
            "user_goal": request.user_goal,
            "timestamp": request.timestamp or datetime.now(),
            "action": "profession_selection"
        }
        
        print(f"[PROFESSION_SELECTION] {log_data}")
        
        return ProfessionSelectionResponse(
            success=True,
            message="Выбор профессии успешно зарегистрирован",
            session_id=session_id
        )
        
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Ошибка логирования выбора профессии: {str(e)}")

@app.post("/get-job-matches", response_model=JobMatchingResponse)
async def get_job_matches(request: JobMatchingRequest):
    """
    Получение подходящих вакансий и логирование запроса
    """
    try:
        # Логируем запрос подборки вакансий
        log_data = {
            "session_id": request.session_id,
            "profession": request.profession,
            "user_role": request.user_role,
            "user_goal": request.user_goal,
            "job_url": request.job_url,
            "timestamp": datetime.now(),
            "action": "job_matching_request"
        }
        
        print(f"[JOB_MATCHING] {log_data}")
        
        # Мок-данные вакансий (позже заменить на реальную логику)
        jobs_map = {
            "AI/ML инженер": [
                {
                    "id": 1,
                    "title": "Senior ML Engineer",
                    "company": "Яндекс",
                    "location": "Москва",
                    "salary": "300 000 - 450 000 ₽",
                    "type": "Полная занятость",
                    "description": "Разработка и внедрение ML-моделей для поисковых алгоритмов"
                },
                {
                    "id": 2,
                    "title": "AI Research Scientist",
                    "company": "Сбер",
                    "location": "Санкт-Петербург",
                    "salary": "250 000 - 400 000 ₽",
                    "type": "Полная занятость",
                    "description": "Исследования в области искусственного интеллекта и NLP"
                }
            ],
            "DevOps": [
                {
                    "id": 1,
                    "title": "Senior DevOps Engineer",
                    "company": "Тинькофф",
                    "location": "Москва",
                    "salary": "250 000 - 350 000 ₽",
                    "type": "Полная занятость",
                    "description": "Автоматизация CI/CD процессов и управление облачной инфраструктурой"
                }
            ]
        }
        
        # Получаем вакансии для профессии
        jobs = jobs_map.get(request.profession, [
            {
                "id": 1,
                "title": "IT Специалист",
                "company": "TechCorp",
                "location": "Москва",
                "salary": "150 000 - 250 000 ₽",
                "type": "Полная занятость",
                "description": "Работа в сфере информационных технологий"
            }
        ])
        
        # Адаптируем вакансии под роль пользователя
        if request.user_role == "student":
            jobs = [{
                **job,
                "title": job["title"].replace("Senior", "Junior") if "Senior" in job["title"] else job["title"],
                "description": job["description"] + " (Позиция для начинающих специалистов)"
            } for job in jobs]
        elif request.user_role == "professional":
            jobs = [{
                **job,
                "title": "Senior " + job["title"] if "Senior" not in job["title"] and "Lead" not in job["title"] else job["title"],
                "description": job["description"] + " (Позиция для опытных специалистов)"
            } for job in jobs]
        
        # Логируем результат
        result_log = {
            "session_id": request.session_id,
            "jobs_count": len(jobs),
            "timestamp": datetime.now(),
            "action": "job_matching_result"
        }
        
        print(f"[JOB_MATCHING_RESULT] {result_log}")
        
        return JobMatchingResponse(
            jobs=jobs,
            total_count=len(jobs),
            session_id=request.session_id
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка получения вакансий: {str(e)}")

@app.post("/upload-resume", response_model=ResumeUploadResponse)
async def upload_resume(file: UploadFile = File(...), session_id: str = None, db: Session = Depends(get_db)):
    """
    Загрузка и извлечение текста из файла резюме с улучшенной диагностикой
    """
    print(f"📄 Получен файл: {file.filename}, тип: {file.content_type}, размер: {file.size if hasattr(file, 'size') else 'неизвестен'}")
    
    try:
        # Расширенный список поддерживаемых типов файлов
        allowed_types = {
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'application/msword': 'doc',
            'text/plain': 'txt',
            'text/plain; charset=utf-8': 'txt',
            'application/octet-stream': 'unknown'  # Для файлов с неопределенным типом
        }
        
        # Определяем тип файла по расширению, если MIME-тип неизвестен
        file_type = None
        if file.content_type in allowed_types:
            file_type = allowed_types[file.content_type]
        elif file.filename:
            extension = file.filename.lower().split('.')[-1]
            if extension == 'pdf':
                file_type = 'pdf'
            elif extension in ['docx']:
                file_type = 'docx'
            elif extension in ['doc']:
                file_type = 'doc'
            elif extension in ['txt']:
                file_type = 'txt'
        
        if not file_type or file_type == 'unknown':
            raise HTTPException(
                status_code=400, 
                detail=f"Неподдерживаемый тип файла: {file.content_type} ({file.filename}). Поддерживаются: PDF, DOCX, DOC, TXT"
            )
        
        # Читаем содержимое файла
        file_content = await file.read()
        actual_size = len(file_content)
        print(f"📊 Фактический размер файла: {actual_size} байт")
        
        if actual_size == 0:
            raise HTTPException(status_code=400, detail="Файл пустой")
        
        # Ограничиваем размер файла (1MB)
        max_size = 1 * 1024 * 1024  # 1MB
        if actual_size > max_size:
            raise HTTPException(status_code=400, detail=f"Файл слишком большой ({actual_size} байт, максимум 1MB)")
        
        # Извлекаем текст в зависимости от типа файла
        print(f"🔍 Начинаем извлечение текста из {file_type.upper()} файла...")
        
        try:
            if file_type == 'pdf':
                extracted_text = extract_text_from_pdf(file_content)
            elif file_type == 'docx':
                extracted_text = extract_text_from_docx(file_content)
            elif file_type == 'doc':
                # DOC файлы сложнее обрабатывать, предлагаем конвертировать в DOCX
                raise ValueError("Файлы .doc не поддерживаются. Пожалуйста, сохраните файл в формате .docx или .txt")
            elif file_type == 'txt':
                extracted_text = extract_text_from_txt(file_content)
            else:
                raise ValueError(f"Неподдерживаемый тип файла: {file_type}")
            
            # Проверяем минимальную длину текста
            text_length = len(extracted_text.strip())
            print(f"📝 Извлечено {text_length} символов текста")
            
            if text_length < 50:
                raise ValueError(f"Извлеченный текст слишком короткий для анализа ({text_length} символов, минимум 50)")
            
            # Показываем первые 200 символов для диагностики
            preview = extracted_text[:200] + "..." if len(extracted_text) > 200 else extracted_text
            print(f"📖 Превью текста: {preview}")
            
            # Сохраняем резюме в базу данных
            try:
                if session_id:
                    # Находим сессию
                    db_session = db.query(DBSession).filter(DBSession.session_id == session_id).first()
                    if db_session:
                        # 🔒 Шифруем текст резюме перед сохранением
                        encrypted_text = encrypt_resume_for_storage(extracted_text)
                        
                        # Создаем запись о резюме
                        resume = Resume(
                            session_id=db_session.id,
                            filename=file.filename,
                            file_type=file_type,
                            file_size=actual_size,
                            extracted_text=encrypted_text  # Сохраняем зашифрованный текст
                        )
                        db.add(resume)
                        db.commit()
                        print(f"💾 Резюме сохранено в базу данных для сессии {session_id}")
            except Exception as db_error:
                print(f"⚠️ Ошибка сохранения в БД: {str(db_error)}")
                # Не прерываем выполнение, если БД недоступна
            
            return ResumeUploadResponse(
                success=True,
                extracted_text=extracted_text,
                file_type=file_type,
                message=f"✅ Текст успешно извлечен из {file_type.upper()} файла ({text_length} символов)"
            )
            
        except ValueError as ve:
            print(f"❌ Ошибка извлечения текста: {str(ve)}")
            return ResumeUploadResponse(
                success=False,
                extracted_text="",
                file_type=file_type,
                message=str(ve)
            )
            
    except HTTPException:
        raise
    except Exception as e:
        print(f"💥 Неожиданная ошибка: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Ошибка обработки файла: {str(e)}")

@app.post("/analyze-resume-ai", response_model=ResumeAnalysisAIResponse)
async def analyze_resume_ai(request: ResumeAnalysisAIRequest, db: Session = Depends(get_db)):
    """
    Анализ резюме с помощью OpenAI API
    """
    try:
        # Проверяем наличие API ключа
        api_key = os.getenv('OPENAI_API_KEY')
        if not api_key:
            return ResumeAnalysisAIResponse(
                analysis="",
                success=False,
                message="OpenAI API ключ не настроен на сервере"
            )
        
        # 🔒 ЗАЩИТА ПЕРСОНАЛЬНЫХ ДАННЫХ
        # Анонимизируем персональные данные перед отправкой в OpenAI
        print(f"🔒 Анонимизация персональных данных...")
        anonymized_text, encrypted_replacements = anonymize_for_ai(request.resume_text)
        
        # Ограничиваем длину анонимизированного текста
        max_resume_length = 6000
        truncated_resume_text = anonymized_text[:max_resume_length]
        if len(anonymized_text) > max_resume_length:
            truncated_resume_text += "\n\n[Текст резюме обрезан для анализа]"
        
        print(f"✅ Персональные данные защищены. Длина анонимизированного текста: {len(truncated_resume_text)} символов")
        
        # Формируем промпт
        prompt = f"""Ты HR-специалист. Проанализируй резюме на позицию "{request.profession}"{f' (вакансия: {request.job_url})' if request.job_url else ''}.

Оцени:
1. Структуру и читаемость
2. Полноту информации и достижения
3. Формулировки (глаголы vs существительные)
4. Карьерный путь
5. Соответствие позиции
6. ATS-оптимизацию

Ответ структурируй:
- Общее впечатление (2-3 предложения)
- Сильные стороны (3-4 пункта)
- Области улучшения (3-4 рекомендации)
- Примеры переформулировок (1-2)
- Адаптация под рынок РФ
- Приоритетные изменения

Резюме:
{truncated_resume_text}"""
        
        # Настраиваем OpenAI клиент с детальной обработкой ошибок
        try:
            print(f"🔑 Инициализация OpenAI клиента...")
            # Инициализируем клиент только с необходимыми параметрами
            # Создаем параметры для OpenAI клиента, исключая proxies
            client_params = {
                "api_key": api_key,
                "timeout": 30.0,  # Таймаут 30 секунд
                "max_retries": 2   # Максимум 2 попытки
            }
            
            # Убираем любые proxy-связанные параметры из окружения
            # для предотвращения автоматической передачи в httpx
            original_env = {}
            proxy_vars = ['HTTP_PROXY', 'HTTPS_PROXY', 'http_proxy', 'https_proxy', 'ALL_PROXY', 'all_proxy']
            for var in proxy_vars:
                if var in os.environ:
                    original_env[var] = os.environ[var]
                    del os.environ[var]
            
            try:
                client = OpenAI(**client_params)
            finally:
                # Восстанавливаем переменные окружения
                for var, value in original_env.items():
                    os.environ[var] = value
            print(f"✅ OpenAI клиент успешно инициализирован")
        except Exception as client_error:
            print(f"❌ Ошибка инициализации OpenAI клиента: {str(client_error)}")
            return ResumeAnalysisAIResponse(
                analysis="",
                success=False,
                message=f"Ошибка инициализации OpenAI клиента: {str(client_error)}"
            )
        
        # Делаем запрос к OpenAI
        try:
            print(f"🤖 Отправка запроса к OpenAI API...")
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "user", "content": prompt}
                ],
                max_tokens=4000,
                temperature=0.7
            )
            print(f"✅ Получен ответ от OpenAI API")
        except Exception as api_error:
            print(f"❌ Ошибка API запроса: {str(api_error)}")
            return ResumeAnalysisAIResponse(
                analysis="",
                success=False,
                message=f"Ошибка API запроса: {str(api_error)}"
            )
        
        analysis = response.choices[0].message.content
        
        # Очищаем markdown элементы
        clean_analysis = analysis\
            .replace('#', '')\
            .replace('**', '')\
            .replace('*', '')\
            .replace('`', '')\
            .strip()
        
        # Сохраняем результат анализа в базу данных
        try:
            if hasattr(request, 'session_id') and request.session_id:
                # Находим сессию
                db_session = db.query(DBSession).filter(DBSession.session_id == request.session_id).first()
                if db_session:
                    # Находим резюме для этой сессии
                    resume = db.query(Resume).filter(Resume.session_id == db_session.id).first()
                    if resume:
                        # Создаем запись об анализе
                        analysis_record = Analysis(
                            resume_id=resume.id,
                            analysis_result=clean_analysis,
                            model_used="gpt-4o-mini",
                            tokens_used=response.usage.total_tokens if hasattr(response, 'usage') else None,
                            processing_time_seconds=None  # Можно добавить замер времени
                        )
                        db.add(analysis_record)
                        db.commit()
                        print(f"💾 Результат анализа сохранен в базу данных")
        except Exception as db_error:
            print(f"⚠️ Ошибка сохранения анализа в БД: {str(db_error)}")
            # Не прерываем выполнение, если БД недоступна
        
        return ResumeAnalysisAIResponse(
            analysis=clean_analysis,
            success=True,
            message="Анализ резюме успешно выполнен"
        )
        
    except Exception as e:
        return ResumeAnalysisAIResponse(
            analysis="",
            success=False,
            message=f"Ошибка при анализе резюме: {str(e)}"
        )

@app.post("/save-user-role", response_model=UserRoleResponse)
async def save_user_role(request: UserRoleRequest, db: Session = Depends(get_db)):
    """
    Сохранение роли пользователя (студент/профессионал)
    """
    try:
        # Генерируем уникальный session_id
        session_id = str(uuid.uuid4())
        
        # Создаем или получаем пользователя
        user = db.query(User).filter(User.telegram_id == request.user_id).first()
        if not user:
            user = User(
                telegram_id=request.user_id,
                username=f"user_{request.user_id}",
                data_processing_consent=True
            )
            db.add(user)
            db.commit()
            db.refresh(user)
        
        # Создаем новую сессию с ролью
        db_session = DBSession(
            session_id=session_id,
            user_id=user.id,
            career_stage=request.user_role
        )
        db.add(db_session)
        db.commit()
        
        # Логируем данные
        log_data = {
            "session_id": session_id,
            "user_id": request.user_id,
            "user_role": request.user_role,
            "timestamp": request.timestamp or datetime.now(),
            "action": "user_role_selection"
        }
        
        print(f"[USER_ROLE_SELECTION] {log_data}")
        
        return UserRoleResponse(
            success=True,
            message="Роль пользователя успешно сохранена",
            session_id=session_id
        )
        
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Ошибка сохранения роли пользователя: {str(e)}")

@app.post("/save-user-goal", response_model=UserRoleResponse)
async def save_user_goal(request: dict, db: Session = Depends(get_db)):
    """
    Сохранение цели пользователя в существующую сессию
    """
    try:
        user_id = request.get('user_id')
        user_goal = request.get('user_goal')
        
        if not user_id or not user_goal:
            raise HTTPException(status_code=400, detail="user_id и user_goal обязательны")
        
        # Находим пользователя
        user = db.query(User).filter(User.telegram_id == user_id).first()
        if not user:
            raise HTTPException(status_code=404, detail="Пользователь не найден")
        
        # Находим последнюю сессию пользователя
        db_session = db.query(DBSession).filter(
            DBSession.user_id == user.id
        ).order_by(DBSession.created_at.desc()).first()
        
        if not db_session:
            raise HTTPException(status_code=404, detail="Сессия не найдена")
        
        # Обновляем цель в сессии
        db_session.career_goal = user_goal
        db.commit()
        
        # Логируем данные
        log_data = {
            "session_id": db_session.session_id,
            "user_id": user_id,
            "user_goal": user_goal,
            "timestamp": datetime.now(),
            "action": "user_goal_selection"
        }
        
        print(f"[USER_GOAL_SELECTION] {log_data}")
        
        return UserRoleResponse(
            success=True,
            message="Цель пользователя успешно сохранена",
            session_id=db_session.session_id
        )
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Ошибка сохранения цели пользователя: {str(e)}")

@app.post("/get-session-state", response_model=SessionStateResponse)
async def get_session_state(request: SessionStateRequest, db: Session = Depends(get_db)):
    """
    Получение состояния сессии пользователя для восстановления экрана
    """
    try:
        # Находим пользователя
        user = db.query(User).filter(User.telegram_id == request.user_id).first()
        if not user:
            return SessionStateResponse(
                success=True,
                has_role=False,
                current_screen="role"
            )
        
        # Находим последнюю сессию пользователя
        last_session = db.query(DBSession).filter(
            DBSession.user_id == user.id
        ).order_by(DBSession.created_at.desc()).first()
        
        if not last_session:
            return SessionStateResponse(
                success=True,
                has_role=False,
                current_screen="role"
            )
        
        # Определяем текущий экран на основе данных сессии
        current_screen = "role"
        if last_session.career_stage:
            current_screen = "goals"
            if last_session.career_goal:
                current_screen = "profession"
                if last_session.selected_profession:
                    current_screen = "main"
        
        return SessionStateResponse(
            success=True,
            has_role=bool(last_session.career_stage),
            user_role=last_session.career_stage,
            selected_profession=last_session.selected_profession,
            career_goal=last_session.career_goal,
            session_id=last_session.session_id,
            current_screen=current_screen
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка получения состояния сессии: {str(e)}")

@app.get("/data-processing-notice")
async def get_data_processing_notice():
    """
    Получение уведомления об обработке персональных данных
    В соответствии с 152-ФЗ "О персональных данных"
    """
    try:
        notice = data_protector.get_data_processing_notice()
        return {
            "success": True,
            "notice": notice,
            "compliance": "152-ФЗ О персональных данных",
            "encryption": "AES-256",
            "anonymization": "Активна"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка получения уведомления: {str(e)}")

@app.get("/admin/stats")
async def get_admin_stats(db: Session = Depends(get_db)):
    """
    Получение статистики для администрирования (только для разработки)
    """
    try:
        # Подсчет пользователей
        total_users = db.query(User).count()
        users_with_consent = db.query(User).filter(User.data_processing_consent == True).count()
        
        # Подсчет сессий
        total_sessions = db.query(DBSession).count()
        sessions_by_profession = db.query(
            DBSession.selected_profession, 
            db.func.count(DBSession.id)
        ).group_by(DBSession.selected_profession).all()
        
        # Подсчет резюме
        total_resumes = db.query(Resume).count()
        resumes_by_type = db.query(
            Resume.file_type,
            db.func.count(Resume.id)
        ).group_by(Resume.file_type).all()
        
        # Подсчет анализов
        total_analyses = db.query(Analysis).count()
        
        return {
            "users": {
                "total": total_users,
                "with_consent": users_with_consent
            },
            "sessions": {
                "total": total_sessions,
                "by_profession": dict(sessions_by_profession)
            },
            "resumes": {
                "total": total_resumes,
                "by_type": dict(resumes_by_type)
            },
            "analyses": {
                "total": total_analyses
            }
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка получения статистики: {str(e)}")



if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)